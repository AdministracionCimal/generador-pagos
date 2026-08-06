"""Convierte MANUAL_USUARIO.md a un .docx para repartir.

Uso:  python tools/manual_a_docx.py [salida.docx]

Cubre lo que usa el manual: títulos, párrafos con **negrita** / `código` /
[links], viñetas, tablas con encabezado, citas (`>`), bloques de código y
separadores. Genera además un índice con las secciones al principio.

Requiere python-docx (no está en requirements: es una herramienta, no parte de
la app). Instalar con:  pip install python-docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
FUENTE = RAIZ / "MANUAL_USUARIO.md"

AZUL = RGBColor(0x1E, 0x3A, 0x5F)      # el mismo de los reportes exportados
GRIS_CODIGO = "F2F2F2"
GRIS_ENCABEZADO = "1E3A5F"

# **negrita**, *cursiva*, `código` y [texto](link) en una sola pasada
_TOKENS = re.compile(
    r"\*\*(?P<negrita>[^*]+)\*\*"
    r"|(?<!\*)\*(?P<cursiva>[^*]+)\*(?!\*)"
    r"|`(?P<codigo>[^`]+)`"
    r"|\[(?P<texto>[^\]]+)\]\((?P<url>[^)]+)\)"
)


def _sombrear(celda, color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")        # nunca "solid": sale negro
    shd.set(qn("w:fill"), color)
    celda._tc.get_or_add_tcPr().append(shd)


def _escribir_inline(parrafo, texto: str) -> None:
    """Agrega el texto al párrafo interpretando el formato de markdown."""
    pos = 0
    for m in _TOKENS.finditer(texto):
        if m.start() > pos:
            parrafo.add_run(texto[pos:m.start()])
        if m.group("negrita"):
            parrafo.add_run(m.group("negrita")).bold = True
        elif m.group("cursiva"):
            parrafo.add_run(m.group("cursiva")).italic = True
        elif m.group("codigo"):
            run = parrafo.add_run(m.group("codigo"))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif m.group("texto"):
            # El link a un archivo del repo no sirve en Word: queda el texto.
            url = m.group("url")
            run = parrafo.add_run(m.group("texto"))
            if url.startswith("http"):
                run.font.color.rgb = AZUL
                run.underline = True
        pos = m.end()
    if pos < len(texto):
        parrafo.add_run(texto[pos:])


def _separador(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    borde = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "BFBFBF")
    borde.append(bottom)
    pPr.append(borde)


def _tabla(doc: Document, filas: list[list[str]]) -> None:
    encabezado, *cuerpo = filas
    tabla = doc.add_table(rows=1, cols=len(encabezado))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = True

    for celda, texto in zip(tabla.rows[0].cells, encabezado):
        celda.text = ""
        p = celda.paragraphs[0]
        _escribir_inline(p, texto)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)
        _sombrear(celda, GRIS_ENCABEZADO)

    for fila in cuerpo:
        celdas = tabla.add_row().cells
        for celda, texto in zip(celdas, fila):
            celda.text = ""
            p = celda.paragraphs[0]
            _escribir_inline(p, texto)
            for run in p.runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()


def _bloque_codigo(doc: Document, lineas: list[str]) -> None:
    tabla = doc.add_table(rows=1, cols=1)
    tabla.style = "Table Grid"
    celda = tabla.rows[0].cells[0]
    _sombrear(celda, GRIS_CODIGO)
    celda.text = ""
    for i, linea in enumerate(lineas):
        p = celda.paragraphs[0] if i == 0 else celda.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(linea)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()


def _fila_de_tabla(linea: str) -> list[str]:
    return [c.strip() for c in linea.strip().strip("|").split("|")]


def _es_separador_de_tabla(linea: str) -> bool:
    return bool(re.match(r"^\|[\s:|-]+\|$", linea.strip()))


def convertir(md: str, doc: Document) -> dict:
    lineas = md.splitlines()
    stats = {"titulos": 0, "tablas": 0, "parrafos": 0, "codigo": 0, "vinetas": 0}
    i = 0
    # El índice del markdown se reemplaza por el que armamos nosotros.
    saltar_indice = False

    while i < len(lineas):
        linea = lineas[i]
        crudo = linea.strip()

        if crudo.startswith("## Índice"):
            saltar_indice = True
            i += 1
            continue
        if saltar_indice:
            if crudo.startswith("---"):
                saltar_indice = False
            i += 1
            continue

        if not crudo:
            i += 1
            continue

        # Títulos
        m_titulo = re.match(r"^(#{1,4})\s+(.*)$", crudo)
        if m_titulo:
            nivel = len(m_titulo.group(1))
            texto = m_titulo.group(2).replace("**", "")
            if nivel == 1:
                doc.add_heading(texto, level=0)
            else:
                doc.add_heading(texto, level=nivel - 1)
            stats["titulos"] += 1
            i += 1
            continue

        # Separador
        if re.match(r"^-{3,}$", crudo):
            _separador(doc)
            i += 1
            continue

        # Bloque de código
        if crudo.startswith("```"):
            i += 1
            bloque = []
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                bloque.append(lineas[i])
                i += 1
            i += 1
            _bloque_codigo(doc, bloque)
            stats["codigo"] += 1
            continue

        # Tabla
        if crudo.startswith("|") and i + 1 < len(lineas) and _es_separador_de_tabla(lineas[i + 1]):
            filas = [_fila_de_tabla(crudo)]
            i += 2
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                filas.append(_fila_de_tabla(lineas[i]))
                i += 1
            _tabla(doc, filas)
            stats["tablas"] += 1
            continue

        # Cita
        if crudo.startswith(">"):
            partes = []
            while i < len(lineas) and lineas[i].strip().startswith(">"):
                partes.append(lineas[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _escribir_inline(p, " ".join(partes))
            for run in p.runs:
                run.italic = True
            stats["parrafos"] += 1
            continue

        # Viñeta o numeral
        m_lista = re.match(r"^([-*]|\d+\.)\s+(.*)$", crudo)
        if m_lista:
            estilo = "List Bullet" if m_lista.group(1) in "-*" else "List Number"
            texto = m_lista.group(2)
            # Continuaciones indentadas de la misma viñeta
            i += 1
            while i < len(lineas) and lineas[i].startswith(("  ", "\t")) and lineas[i].strip() \
                    and not re.match(r"^\s*([-*]|\d+\.)\s+", lineas[i]):
                texto += " " + lineas[i].strip()
                i += 1
            p = doc.add_paragraph(style=estilo)
            _escribir_inline(p, texto)
            stats["vinetas"] += 1
            continue

        # Párrafo (junta las líneas siguientes hasta un corte)
        partes = [crudo]
        i += 1
        while i < len(lineas):
            siguiente = lineas[i].strip()
            if (not siguiente or siguiente.startswith(("#", "|", ">", "```", "---"))
                    or re.match(r"^([-*]|\d+\.)\s+", siguiente)):
                break
            partes.append(siguiente)
            i += 1
        p = doc.add_paragraph()
        _escribir_inline(p, " ".join(partes))
        stats["parrafos"] += 1

    return stats


def _indice(doc: Document, md: str) -> None:
    doc.add_heading("Contenido", level=1)
    for linea in md.splitlines():
        if linea.startswith("## ") and "Índice" not in linea:
            texto = linea[3:].strip().replace("**", "")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            _escribir_inline(p, texto)
        elif linea.startswith("### "):
            texto = linea[4:].strip().replace("**", "")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.55)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(texto)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_page_break()


def main() -> int:
    salida = Path(sys.argv[1]) if len(sys.argv) > 1 else RAIZ / "MANUAL_USUARIO.docx"
    md = FUENTE.read_text(encoding="utf-8")

    doc = Document()
    seccion = doc.sections[0]
    seccion.page_width, seccion.page_height = Inches(8.5), Inches(11)   # Carta
    for lado in ("left_margin", "right_margin"):
        setattr(seccion, lado, Inches(0.8))
    seccion.top_margin = seccion.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    # Portada mínima + índice antes del cuerpo
    titulo = doc.add_heading("Generador de Pagos", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Manual de usuario")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = AZUL
    doc.add_page_break()
    _indice(doc, md)

    # El H1 del markdown ya está en la portada
    cuerpo = re.sub(r"^# .*\n", "", md, count=1)
    stats = convertir(cuerpo, doc)

    doc.save(salida)
    print(f"{salida.name}: {stats}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
